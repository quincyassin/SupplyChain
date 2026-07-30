#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成《分单宝-产品功能介绍》Word 文档。@author huangxinsong"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os


def set_run_font(run, size=12, bold=False, color=None, font_name="微软雅黑"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    if color:
        run.font.color.rgb = color


def build_document():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    def add_heading_cn(text, level=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text)
        sizes = {1: 18, 2: 14, 3: 12}
        set_run_font(run, size=sizes.get(level, 12), bold=True, color=RGBColor(0x1A, 0x1A, 0x1A))
        return p

    def add_para(text, size=11, space_after=6, first_indent=True):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.35
        if first_indent:
            p.paragraph_format.first_line_indent = Cm(0.74)
        run = p.add_run(text)
        set_run_font(run, size=size)
        return p

    def add_bullet(text, size=11):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.3
        p.clear()
        run = p.add_run(text)
        set_run_font(run, size=size)
        return p

    def add_feature_block(title, desc, bullets=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(title)
        set_run_font(run, size=12, bold=True, color=RGBColor(0x2C, 0x3E, 0x50))
        add_para(desc, first_indent=False)
        if bullets:
            for item in bullets:
                add_bullet(item)

    # 封面
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(40)
    title.paragraph_format.space_after = Pt(8)
    set_run_font(title.add_run("分单宝"), size=28, bold=True, color=RGBColor(0x1A, 0x1A, 0x1A))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(6)
    set_run_font(subtitle.add_run("产品功能介绍"), size=18, bold=True, color=RGBColor(0x44, 0x44, 0x44))

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline.paragraph_format.space_after = Pt(28)
    set_run_font(
        tagline.add_run("多平台订单导入 · 智能分单 · 批量回单 · 对账售后一站完成"),
        size=11,
        color=RGBColor(0x66, 0x66, 0x66),
    )

    # 一、产品概述
    add_heading_cn("一、产品概述", 1)
    add_para(
        "分单宝是一款面向电商运营、代发与分销场景的订单处理工具。"
        "支持将淘宝、拼多多等多平台导出的 Excel 订单统一导入系统，"
        "按商家规则自动分单，完成物流回单、对账导出与售后跟踪，形成从接单到结账的完整业务闭环。"
    )
    add_para(
        "产品以本地 Web 应用形式提供，Mac / Windows 单机版双击即可使用，数据保存在本机，"
        "无需额外安装 Java 或数据库环境，适合小团队与个人运营日常使用。"
    )

    # 二、核心价值
    add_heading_cn("二、核心价值", 1)
    values = [
        ("多平台兼容", "各电商平台 Excel 表头不同，系统通过预配置的平台表头映射与智能列识别，上传即可自动归类入库，减少手工整理。"),
        ("智能分单", "依据商家配置的商品关键字，将订单自动分配给对应供货商家；未识别订单归入「未定义」，补充规则后可一键重新分单。"),
        ("全流程闭环", "覆盖导入、分单导出、批量回单、回单导出回平台、商家/平台对账等关键环节，降低跨表格来回切换成本。"),
        ("经营数据可见", "维护商品成本价与供货价后，首页按日汇总营业额、成本与利润，便于快速掌握经营情况。"),
        ("售后与数据治理", "支持售后标记、完结与导出；误删订单可进回收站恢复；历史订单可归档，保持日常列表清爽可控。"),
        ("零门槛部署", "单机版内置运行环境，解压或安装后即可在浏览器中使用；数据目录可备份迁移，便于更换电脑。"),
    ]
    for i, (name, desc) in enumerate(values, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.35
        set_run_font(p.add_run(f"{i}. {name}："), size=11, bold=True)
        set_run_font(p.add_run(desc), size=11)

    # 三、功能模块详解
    add_heading_cn("三、功能模块详解", 1)
    add_para(
        "系统主界面包含：首页、对账、售后、回收站、商品价格维护、系统配置、使用手册等模块，以下按模块说明能力要点。",
        first_indent=False,
    )

    add_heading_cn("3.1 首页（订单处理）", 2)
    add_para(
        "首页是日常订单工作台，承担上传、浏览、分单、回单、导出与售后标记等核心操作。",
        first_indent=False,
    )
    add_feature_block(
        "上传订单 Excel",
        "选择平台导出的订单表格上传后，系统按已配置的表头映射自动识别平台并入库；"
        "无法匹配商家规则的订单会进入「未定义」，可在商家配置中补充识别词后重新分单。",
    )
    add_feature_block(
        "分单日期与经营汇总",
        "通过顶部日期范围与左侧分单日期列表切换查看历史订单（不可选未来日期）。"
        "页面底部展示所选日期的订单数量、营业额、成本价与利润。",
    )
    add_feature_block(
        "搜索与筛选",
        "支持多条件快速定位订单：",
        [
            "关键词搜索：商家名、平台名、系统编号、快递单号、订单编号等；",
            "回单状态：全部 / 未回单 / 已回单；",
            "平台、商家 Tab：在表格上方进一步缩小查看范围。",
        ],
    )
    add_feature_block(
        "按商家分单",
        "将当前日期范围内的订单按规则分配给各商家。已正确分单的订单不会被改动；"
        "未分或需重分的订单会重新识别。确认后按商家生成 Excel，保存位置由「导出配置」决定（本地文件夹或浏览器下载）。",
    )
    add_feature_block(
        "填写物流信息（回单）",
        "商家发货后，可批量粘贴或录入系统编号、快递公司、快递单号等信息，订单标记为「已回单」。"
        "查找范围以顶部「分单日期」为准，不受当前平台/商家 Tab 影响。",
    )
    add_feature_block(
        "回单导出与所选导出",
        "回单导出：将已回单订单按平台导出为表格，便于回传到各电商平台。"
        "所选数据导出：将勾选订单导出为文件，用于备份或二次整理。",
    )
    add_feature_block(
        "行内操作",
        "每笔订单支持查看详情、标记售后（填写原因后进入售后页）、取消售后、删除等操作；支持批量删除已勾选订单。",
    )

    add_heading_cn("3.2 对账", 2)
    add_para(
        "按商家或平台统计指定分单日期范围内的订单，并导出对账 Excel，便于与商家或平台侧数据核对。",
        first_indent=False,
    )
    add_feature_block(
        "商家对账",
        "选定日期范围后，在商家列表中选择目标商家（列表显示订单笔数），一键导出该商家对账表。"
        "「未定义」及尚未分单的商家不会出现在列表中。",
    )
    add_feature_block(
        "平台对账",
        "切换至「平台对账」，选择平台后导出对账表格，用于与平台账单或后台数据对照。",
    )
    add_feature_block(
        "导出保存规则",
        "对账文件与分单、回单共用同一套导出配置：可保存到指定本地目录（按日期分子文件夹，对账文件置于「对账」目录），"
        "也可通过浏览器下载；保存到本地时，导出完成后会尽量自动打开对应文件夹。",
    )

    add_heading_cn("3.3 售后", 2)
    add_para(
        "集中管理在首页标记为需售后的订单，与首页售后标记联动，形成售后处理台账。",
        first_indent=False,
    )
    add_feature_block(
        "查询与筛选",
        "默认展示近 30 天数据，可按日期范围、售后状态（需售后 / 售后完结）以及系统编号、订单号、快递单号等条件筛选。",
    )
    add_feature_block(
        "处理动作",
        "支持将订单标记为售后完结、取消售后，以及导出售后数据，便于留存与复盘。完结订单在列表中会高亮区分。",
    )

    add_heading_cn("3.4 回收站", 2)
    add_para(
        "首页删除的订单进入回收站（软删除），避免误删造成不可恢复的损失。",
        first_indent=False,
    )
    add_feature_block(
        "恢复与清除",
        "支持按日期范围与关键词查询，可对订单执行批量/单条恢复，或永久清除。"
        "默认查询近 30 天数据。回收站恢复与「数据归档」恢复是两套不同机制，请按场景选用。",
    )

    add_heading_cn("3.5 商品价格维护", 2)
    add_para(
        "维护各平台商品的成本价与供货价，供首页利润计算使用。商品列表来源于当前订单与归档订单的汇总结果。",
        first_indent=False,
    )
    add_feature_block(
        "改价与批量维护",
        "支持在表格内直接修改价格；提供模板下载与 Excel 批量导入；"
        "支持批量删除价格记录（仅删除价格配置，不删除订单本身）。",
    )

    add_heading_cn("3.6 系统配置", 2)
    add_para(
        "系统配置是后台能力中心，决定订单如何识别、如何展示、如何导出以及如何授权使用，包含以下子模块。",
        first_indent=False,
    )
    add_feature_block(
        "表头映射（平台）",
        "配置各电商平台 Excel 列与系统标准字段的对应关系。上传订单时依赖此配置完成平台识别与字段入库。",
    )
    add_feature_block(
        "字段映射",
        "自定义界面字段的显示名称，使列表表头更贴合团队内部叫法，不影响底层数据存储。",
    )
    add_feature_block(
        "商家配置",
        "维护商家名称及其商品关键字识别规则。分单时按关键字匹配商家；识别词越完整，「未定义」订单越少。",
    )
    add_feature_block(
        "导出配置",
        "设置导出文件保存方式：本地目录（可按日期自动建子文件夹）或浏览器下载。单机版支持本机选择文件夹。",
    )
    add_feature_block(
        "数据归档",
        "将历史订单从日常表移入归档表，减轻首页数据量；支持预览与恢复，便于阶段性清理与回溯。",
    )
    add_feature_block(
        "软件授权",
        "单机版采用一机一码离线授权：绑定机器码，通过激活码完成授权与到期控制，保障正版使用。",
    )

    add_heading_cn("3.7 使用手册", 2)
    add_para(
        "软件内置按模块分节的操作说明，内容覆盖首页、对账、售后、配置等日常场景，可随时在界面中查阅，降低上手成本。",
        first_indent=False,
    )

    # 四、典型业务流程
    add_heading_cn("四、典型业务流程", 1)
    add_para("以日常代发/分销作业为例，推荐按以下顺序使用各模块：", first_indent=False)
    steps = [
        ("初次配置", "在「系统配置」中完成平台表头映射、商家关键字、导出目录（及单机版授权激活）。"),
        ("导入订单", "从各电商平台导出订单 Excel，在首页「上传订单 Excel」导入系统。"),
        ("智能分单", "确认分单日期范围后，执行「按商家分单」，将订单分配给各供货商家并导出商家表格。"),
        ("商家发货与回单", "商家发货后，在首页「填写物流信息」批量回填快递信息；再执行「回单导出」，将结果回传到对应平台。"),
        ("价格与利润", "在「商品价格维护」中维护成本/供货价，首页即可查看营业额、成本与利润汇总。"),
        ("对账结算", "在「对账」中按商家或平台导出对账表，完成阶段性结算核对。"),
        ("售后处理", "异常订单在首页标记售后，于「售后」页跟进完结或取消，必要时导出售后明细。"),
        ("数据治理", "误删订单可从「回收站」恢复；历史订单可通过「数据归档」移出日常列表，需要时再恢复。"),
    ]
    for i, (name, desc) in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.35
        set_run_font(p.add_run(f"步骤 {i}｜{name}："), size=11, bold=True)
        set_run_font(p.add_run(desc), size=11)

    end = doc.add_paragraph()
    end.paragraph_format.space_before = Pt(24)
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(end.add_run("— 文档结束 —"), size=10, color=RGBColor(0x99, 0x99, 0x99))
    return doc


def main():
    root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    out_path = os.path.join(root, "docs", "分单宝-产品功能介绍.docx")
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc = build_document()
    doc.save(out_path)
    print(f"OK: {out_path}")
    print(f"size: {os.path.getsize(out_path)}")


if __name__ == "__main__":
    main()
